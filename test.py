# coding=utf-8
import sys
from openpyxl import Workbook
from openpyxl import load_workbook
import os
from docx import Document
from docx.shared import Inches
import logging

def str_len(str):
    try:
        row_l=len(str)
        utf8_l=len(str.encode('utf-8'))
        return (utf8_l-row_l)/2+row_l
    except:
        return None
    return None
	
def replcae_str(src, dest):
	if (len(src.strip()) > str_len(dest)) :
		length = len(src) - len(src.strip()) + str_len(dest)
	else:
		dest_len = str_len(dest) - len(dest)
		length = len(src) - dest_len
		#length = len(src)
	
	return src.rstrip().replace(src.strip(),dest).ljust(int(length))

def set_paragraphs_text(pa, key, value):
	for p in pa:
		for runitem in p.runs:
			src = runitem.text
			#length = len(key)
			#src_cmp = src.strip()[0:length]
			#print(src.strip())
			if (src.strip() == key):
				#print('find it')
				#print(src)
				runitem.text = replcae_str(runitem.text, value)
			#print(src)

def set_table_text(table, key, value):
	for c in table:
		cells = c._cells
		for cell in cells:
			set_paragraphs_text(cell.paragraphs, key, value)
			set_table_text(cell.tables, key, value)

def set_item_text(doc, key, value):
	pas = doc.paragraphs
	set_paragraphs_text(doc.paragraphs, key, value)
	set_table_text(doc.tables, key, value)

def get_item_text(sheet, r, c):
	cell_val = sheet.cell(row = r, column = c).value
	return cell_val

def endWith(s,*endstring):
	array = map(s.endswith,endstring)
	if True in array:
		return True
	else:
		return False

def replace_file(filename, items):
	#try:
	document = Document(filename)
	for key,value in items.items():
		#print(key)
		#print(value)
		set_item_text(document, key, value)
	document.save('结果/' + filename)
	#except:
	#	print("can not open the file")

def format_data(dt):
	year = dt.strftime('%Y')
	month = dt.strftime('%m')
	month = str(int(month))
	day = dt.strftime('%d')
	return year + '年' + month + '月' + day + '日'

def CleanDir(Dir):
	if os.path.isdir(Dir):
		paths = os.listdir(Dir)
		for path in paths:
			filePath = os.path.join(Dir, path)
			if os.path.isfile(filePath):
				try:
					os.remove(filePath)
				except os.error:
					autoRun.exception( "remove %s error." %filePath )#引入logging
			elif os.path.isdir(filePath):
				if filePath[-4:].lower() == ".svn".lower():
					continue
				shutil.rmtree(filePath,True)
	return True
	
def load_xml():
	dist_ret = {}
	xml_name = ' '
	s = os.listdir('./')
	for i in s:
		if (i[0] == '~'):
			continue
		if (endWith(i, '.xlsx')):
			xml_name = i
			break
	if (xml_name == ' '):
		write_log('找不到xml模板')

	try:
		bk = load_workbook(xml_name, read_only = True,data_only = True)
		sh = bk["MySheet"]
	except:
		write_log("no sheet named MySheet in file " + xml_name)

	nrows = sh.max_row
	for i in sh.iter_rows():
		key = i[1].value
		cell = i[3]
		value = cell.value
		write_log("key is " + str(key) + ",format is " + str(cell.number_format) + ", value is " + str(value))
		
		if (cell.number_format == '0.00%'):
			value = value * 100
			value = str(value) + '%'
	
		if (cell.number_format == '0%'):
			value = value * 100
			value = str(int(value)) + '%'
		
		if (cell.number_format == 'mm-dd-yy'):
			value = format_data(value)
		
		dist_ret[key] = str(value)
	return dist_ret

if (os.path.exists("结果")):
	CleanDir("结果")
else:
	os.mkdir("结果")
	
#init logging
logger = logging.getLogger("log_transfer")
logger.setLevel(logging.INFO)

log_fh = logging.FileHandler("./结果/log")
log_fh.setLevel(logging.INFO)

fmt = "%(asctime)-15s %(levelname)s %(filename)s %(lineno)d %(process)d %(message)s"
datefmt = "%a %d %b %Y %H:%M:%S"
formatter = logging.Formatter(fmt, datefmt)

log_fh.setFormatter(formatter)
logger.addHandler(log_fh)

def write_log(str):
	global logger
	logger.info(str)

key_dist = load_xml()


s = os.listdir('./')
for i in s:
	if (endWith(i, '.docx')):
		write_log("transfer file:" + i)
		replace_file(i, key_dist)